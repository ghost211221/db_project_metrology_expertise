# Converts a docx file with tables and images to (simple) HTML
# ORIGINAL CODE BY DAVID SSALI AT SOURCE: https://medium.com/@dvdssali/docx-to-html-1374eb6491a1
#
# Requires the Python module 'python-docx' <https://python-docx.readthedocs.io>
#
# Example use:
# >>> s = convert('./SOURCEDIR', 'SAMPLE.docx')
# >>> print(s)
# or
# $ python .\convert-docx-to-html.py > temp.html

from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import xml.etree.ElementTree as ET
import time
import re
import os

class Docx2HtmlConverter():

    def __init__(self):

        self.document = None

        self.body = '<div class="page">'

        self.list_items = []

        self.page = 0

    def convert(self, dir_path,file):
        """
            dir_path = "./SOURCEDIR"
            file = "SAMPLE.docx"

            image_path will be a directory to store image files; if the directory does not exist it will be created
        """

        splitname = re.split('[. ]', file)
        if len(splitname) > 2:
            image_path=  ''.join(splitname[:2]) + '-images'
        else:
            image_path= splitname[0] + '-images'

        if not os.path.isdir(image_path):
            os.mkdir(image_path)

        self.document = Document(dir_path + '/' + file)

        self.__covertDocument()

    def getHtml(self):

        return self.body

    def __covertDocument(self):
        pass

    def __detectPageBreak(self, block, type='paragraph'):
        pass

    def __




def convert(dir_path,file):
    ##dir_path = "./SOURCEDIR"
    ##file = "SAMPLE.docx"
    # image_path will be a directory to store image files; if the directory does not exist it will be created
    splitname = re.split('[. ]', file)
    if len(splitname) > 2:
        image_path=  ''.join(splitname[:2]) + '-images'
    else:
        image_path= splitname[0] + '-images'
    if not os.path.isdir(image_path):
        os.mkdir(image_path)
    document = Document(dir_path + '/' + file)
    body = '<div class="page">'
    list_items = []
    pn = 1
    for block in iter_block_items(document):

        if isinstance(block, Paragraph):
            for run in block.runs:
                if 'lastRenderedPageBreak' in run._element.xml:
                    body += '</div>\n<div class="page">'
                    pn+=1
                elif 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                    body += '</div>\n<div class="page">'
                    pn+=1

            tmp_heading_type = get_heading_type(block)
            if re.match("List\sParagraph",tmp_heading_type):
                list_items.append("<li>" + block.text + "</li>\n")
            elif not block.text.strip():
                # переводим пустые строку документа в пустые строки
                list_items.append("<br/>\n")
            else:
                images = render_image(document,block,image_path,image_path)
                if len(list_items) > 0:
                    body += render_list_items(list_items)
                    list_items = []
                if len(images) > 0:
                    body = body + images
                else:
                    # modified to use a different outer_tag if a 'Heading' style is found in the original paragraph
                    if 'Heading' in tmp_heading_type:
                        outer_tag = 'h'+tmp_heading_type.split(' ')[-1]
                    else:
                        outer_tag = 'p'
                    body = body + render_runs(block.runs, outer_tag)
        elif isinstance(block, Table):
            body += render_table(block,document,image_path, pn)

    body += '</div>\n'
     
    return body
   
def iter_block_items(parent):
    """
    Generate a reference to each paragraph and table child within *parent*,
    in document order. Each returned value is an instance of either Table or
    Paragraph. *parent* would most commonly be a reference to a main
    Document object, but also works for a _Cell object, which itself can
    contain paragraphs and tables.
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)
            
def table_print(block):
    table = block
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                print(paragraph.text,'  ',end='')
        # print("\n")
        
# Modified to treat cell content as a set of blocks to process 
def render_table(block,document,image_path, pn):
    table = block
    
    html = "<table class='table table-bordered' border='1'>\n"
    for row in table.rows:
        if 'lastRenderedPageBreak' in row._element.xml:
            html += '</div>\n<div class="page">'
            pn+=1
        elif 'w:br' in row._element.xml:
            html += '</div>\n<div class="page">'
            pn+=1
        html += "<tr>\n"
        for cell in row.cells:
            html += "<td>\n"
            # --- 
            cbody = ""
            clist_items = []
            for cblock in iter_block_items(cell):
                if isinstance(cblock, Paragraph):
                    tmp_heading_type = get_heading_type(cblock)
                    if re.match("List\sParagraph",tmp_heading_type):
                        clist_items.append("<li>" + cblock.text + "</li>\n")
                    else:
                        images = render_image(document,cblock,image_path,image_path)
                        if len(clist_items) > 0:
                            cbody += render_list_items(clist_items)
                            clist_items = []
                        if len(images) > 0:
                            cbody = cbody + images
                        else:
                            cbody = cbody + render_runs(cblock.runs)
                elif isinstance(cblock, Table):
                    cbody += render_table(cblock, document, image_path)
            html += cbody + " "
            html += "</td>\n"
        html += "</tr>\n"
    html += "</table>\n"
    return html
    
# Modified to use a different outer_tag if a 'Heading' style is found in the original paragraph
def render_runs(runs, outer_tag='p'):
    html = "<" + outer_tag + ">"
    for run in runs:
        html = html + run.text
    html += "</" + outer_tag + ">\n"
    return html
    
def render_list_items(items):
    html = "<ul>\n"
    for item in items:
        html += item
    html += "</ul>\n"
    return html
    
def get_heading_type(block):
    return block.style.name
    
def render_image(document,par,dir_path,book_id):
    """get all of the images in a paragraph 
    :param par: a paragraph object from docx
    :return: a list of r:embed 
    """
    ids = []
    root = ET.fromstring(par._p.xml)
    namespace = {
             'a':"http://schemas.openxmlformats.org/drawingml/2006/main", \
             'r':"http://schemas.openxmlformats.org/officeDocument/2006/relationships", \
             'wp':"http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"}
    inlines = root.findall('.//wp:inline',namespace)
    for inline in inlines:
        imgs = inline.findall('.//a:blip', namespace)
        for img in imgs:     
            id = img.attrib['{{{0}}}embed'.format(namespace['r'])]
            ids.append(id)
    inlines = root.findall('.//wp:anchor',namespace)
    for inline in inlines:
        imgs = inline.findall('.//a:blip', namespace)
        for img in imgs:     
            id = img.attrib['{{{0}}}embed'.format(namespace['r'])]
            ids.append(id)
    response = ""
    if len(ids) > 0:
        for id in ids:
            image_part = document.part.related_parts[id]
            millis = int(round(time.time() * 1000))
            file_name = str(id) + "-" + str(book_id) + "-" + str(millis) + ".png"
            fr = open(dir_path + "/" + file_name, "wb")
            fr.write(image_part._blob)
            fr.close()
            response += "<img border='1' src='" + dir_path + "/" + file_name + "' class='img-responsive'/>\n"
    return response
    
if __name__ == '__main__':
    s = convert('.', 'init_document.docx')
    with open('text.html', 'w', encoding='utf8') as f:
        f.write(s)