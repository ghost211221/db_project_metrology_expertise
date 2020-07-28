# Converts a docx file with tables and images to (simple) HTML
# ORIGINAL CODE BY DAVID SSALI AT SOURCE: https://medium.com/@dvdssali/docx-to-html-1374eb6491a1
#
# Requires the Python module 'python-docx' <https://python-docx.readthedocs.io>
#
# Example use:
# >>> s = convert('./SOURCEDIR', 'SAMPLE.docx')
# >>> print(s)
# or
# $ python .\convert-docx-to-html.py > temp.html


"""
Конвертер работает. Выдает html документа.
Доработки:
1. сделать выборку стилей параграфов в словарь - из словаря генерить строку css;
2. сделать выборку стилей таблиц в словарь - из словаря генерить строку css;
3. сделать выборку стилей ячеек таблицы в словарь - из словаря генерить строку css;
4. настроить генерацию отступов для вложенных тэгов;
5. перевести определение размеров на pt, как в word.
"""

from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import xml.etree.ElementTree as ET
import time
import re
import os
from base64 import b64encode


class Docx2HtmlConverter():

    def __init__(self):

        self.document = None
        self.file_name = ''
        self.image_path = ''

        self.body_json = []
        self.list_items = []

        # номер страницы, понадобится когда получится сделать нормальное разбиение по страницам
        self.page = 1

        # нумерация основых элементов документа на странице
        self.paragraph = 1
        self.header = 1
        self.ul = 1
        self.img = 1
        self.table = 1

        # аналогично, но для таблицы
        self.t_paragrpah = 1
        self.t_header = 1
        self.t_ul = 1
        self.t_img = 1
        self.t_row = 1
        self.t_cell = 1

        self.li = 1
        self.br = 1

        self.matrix = []


        # сразу добавляем первую страницу
        self.__addNewPage()


    def convert(self, dir_path, file):
        """
            dir_path = "./SOURCEDIR"
            file = "SAMPLE.docx"

            image_path will be a directory to store image files; if the directory does not exist it will be created
        """
        self.file_name = file
        splitname = re.split('[. ]', file)
        if len(splitname) > 2:
            self.image_path =  dir_path + ''.join(splitname[:2]) + '-images'
        else:
            self.image_path = dir_path + splitname[0] + '-images'

        if not os.path.isdir(self.image_path):
            os.mkdir(self.image_path)

        self.document = Document(dir_path + file)

        self.__covertDocument()

    def getJSON(self):
        """ возврат сгенеренного JSON """

        return self.body_json

    def __chkParent(self, root, parent):
        """ проверка то parent в root существует """
        if 'children' in root:
            for child in root['children']:
                if child['id'] == parent:
                    return True

        return False


    def __elemIdxInRoot(self, root, _id):

        if root.get('chidren'):
            for child in root['children']:
                if child['id'] == _id:
                    return root['children'].index(child)

    def __getParentById(self, root, id):
        for item in root:
            if item['id'] == id:
                return item


    def __covertDocument(self):
        for block in self.__iter_block_items(self.document):
            parent = self.__getParentById(self.body_json, f'page-{self.page}')

            if 'w:drawing' in block._element.xml:
                rIdx = re.search(r'rId[0-9]+', block._element.xml)[0]
                blob_ = self.document.inline_shapes.part.related_parts[rIdx]._blob

                self.__addImage(parent, rIdx, blob_)

            elif isinstance(block, Paragraph):
                # self.__detectPageBreakBlock(block)
                tmp_heading_type = self.__get_heading_type(block)

                if re.match(r"List\sParagraph", tmp_heading_type) or 'w:ilvl' in block._element.xml and 'w:numId w:val="2"' in block._element.xml \
                    and 'w:numPr' in block._element.xml and 'w:pPr' in block._element.xml:

                    if parent and not self.__chkParent(parent, f'page-{self.page} ul-{self.ul}'):
                        self.__addNewUlToParent(parent)

                    self.list_items.append(block)

                elif not block.text.strip():
                    # переводим пустые строку документа в пустые строки
                    self.__addBrToParent(parent)

                else:
                    if len(self.list_items) > 0:
                        self.__add_list_items(
                            self.list_items,
                            self.__getParentById(parent['children'], f'page-{self.page} ul-{self.ul}')
                        )
                        
                    else:
                        # modified to use a different outer_tag if a 'Heading' style is found in the original paragraph
                        if 'Heading' in tmp_heading_type:
                            self.__add_heading(parent, block, tmp_heading_type.split(' ')[-1])

                        else:
                            self.__add_paragraph(parent, block)

            elif isinstance(block, Table):
                self.__addNewTable(parent, 'w:tcBorders' in block._element.xml)
                table_parent = self.__getParentById(parent['children'], f'{parent["id"]} table-{self.table}')
                self.__prepare_table(block, table_parent)


    def __iter_block_items(self, parent):
        """
        Generate a reference to each paragraph and table child within *parent*,
        in document order. Each returned value is an instance of either Table or
        Paragraph. *parent* would most commonly be a reference to a main
        Document object, but also works for a _Cell object, which itself can
        contain paragraphs and tables.
        """
        if isinstance(parent, _Document):
            parent_elm = parent.element.body
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        else:
            raise ValueError("something's not right")
        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    def __detectPageBreakBlock(self, block, blockType='paragraph'):
        """ поиск разрыва страницы в блоке """
        if blockType == 'paragraph':
            for run in block.runs:
                self.__detectPageBreak(run, blockType='paragraph')

        elif blockType == 'table':
            self.__detectPageBreak(block, blockType='table')

        else:
            raise Exception(f"""Page break detector
                unexpected block type - {blockType}""")

    def __detectPageBreak(self, block, blockType='paragraph'):
        """ поиск разрыва строки в Paragraph run или в table row"""
        # pass
        if 'lastRenderedPageBreak' in block._element.xml or 'w:br' in block._element.xml and 'type="page"' in block._element.xml:
            self.page += 1
            # self.__addNewPage()
            # self.paragraph = 1
            # if blockType == 'paragraph':
            #     self.body += f'</div>\n<div class="page page-{self.page}" style="width: 975px;">\n'
            # else:
            #     self.body += f'<\ttbody>\n</div>\n<div class="page page-{self.page}" style="width: 975px;">\t<table>\n\t<tbody>\n'
                

    def __get_heading_type(self, block):
        return block.style.name

    def __add_list_items(self, items, parent):
        if parent and 'children' in parent:
            for item in items:
                style = self.__get_align(item)
                style += self.__get_font_size(item)
                style += self.__get_vertical_spacing(item)

                parent['children'].append(
                    {
                        'type': 'li',
                        'class': f'listitem listitem-{self.li}',                
                        'ref': f'{parent["id"]} listitem-{self.li}',
                        'style': style,
                        'text': item.text
                    }
                )

                self.li +=  1

            self.list_items = []
            self.ul += 1
            self.li = 1

    
    def __prepare_table(self, block, root):
        """ Modified to treat cell content as a set of blocks to process  """
        table = block

        self.matrix = []

        for row in table.rows:
            row_ = []

            self.__addNewRow(root, row)

            parent = self.__getParentById(root['children'], f'{root["id"]} row-{self.t_row}')

            for cell in row.cells:
                if '<w:vMerge w:val="restart"/>' in cell._element.xml:
                    row_.append(1)
                elif '<w:vMerge w:val="continue"/>' in cell._element.xml:
                    row_.append(2)
                else:
                    row_.append(0)
                    
                self.__addNewCell(parent, cell, 'w:tcBorders' in cell._element.xml)

                parent_ = self.__getParentById(parent['children'], f'{parent["id"]} cell-{self.t_cell}')

                for cblock in self.__iter_block_items(cell):
                    if 'w:drawing' in cblock._element.xml:
                        rIdx = re.search(r'rId[0-9]+', cblock._element.xml)[0]
                        blob_ = self.document.inline_shapes.part.related_parts[rIdx]._blob

                        self.__addImage(parent_, rIdx, blob_)

                    elif isinstance(cblock, Paragraph):
                        # self.__detectPageBreakBlock(block)
                        tmp_heading_type = self.__get_heading_type(cblock)

                        if re.match(r"List\sParagraph", tmp_heading_type) or 'w:ilvl' in cblock._element.xml and 'w:numId w:val="2"' in cblock._element.xml \
                            and 'w:numPr' in cblock._element.xml and 'w:pPr' in cblock._element.xml:

                            if parent_ and not self.__chkParent(parent_, f'{root["id"]} ul-{self.ul}'):
                                self.__addNewUlToParent(parent_)

                            self.list_items.append(cblock)

                        elif not cblock.text.strip():
                            # переводим пустые строку документа в пустые строки
                            self.__addBrToParent(parent_)

                        else:
                            if len(self.list_items) > 0:
                                self.__add_list_items(
                                    self.list_items,
                                    self.__getParentById(parent_['children'], f'{root["id"]} ul-{self.ul}')
                                )
                                
                            else:
                                # modified to use a different outer_tag if a 'Heading' style is found in the original paragraph
                                if 'Heading' in tmp_heading_type:
                                    self.__add_heading(parent_, cblock, tmp_heading_type.split(' ')[-1])

                                else:
                                    self.__add_paragraph(parent_, cblock)

                    elif isinstance(cblock, Table):
                        self.__addNewTable(parent_, 'w:tcBorders' in cblock._element.xml)
                        table_parent = self.__getParentById(parent_['children'], f'{root["id"]} table-{self.table}')
                        self.__prepare_table(cblock, table_parent)

                self.t_cell += 1

            self.matrix.append(row_)

            self.t_row += 1        

        while self.__matrix_sum() > 0:
            span = self.__get_span()
            if span:
                self.__add_span(root, span)

        self.table += 1

    def __matrix_sum(self):
        sum_ = 0
        for row in self.matrix:
            for cell in row:
                sum_ += cell

        return sum_

    def __get_span(self):
        rown = 0
        for row in self.matrix:
            coln = 0
            for cell in row:
                if cell == 1:
                    span = [rown, coln, 1]
                    row_ = rown

                    while row_ <= len(self.matrix)-1 and (self.matrix[row_][coln] == 1 or self.matrix[row_][coln] == 2):
                        span[2] += 1
                        self.matrix[row_][coln] = 0
                        row_ += 1

                    span[2] -= 1

                    return span

                coln += 1
            rown += 1

    def __add_span(self, root, span):
        for row in root['children']:
            if f'row-{span[0]+1}' in row['id']:
                # нашли строку, ищем ячейку
                for cell in row['children']:
                    if f'cell-{span[1]+1}' in cell['id']:
                        # нашли ячейку
                        cell['rowspan'] = span[2]
                

    def __get_block_style(self, block):
        css = ''

        return css

    def __addNewPage(self):
        self.body_json.append(
            {
                'type': 'div',
                'id': f'page-{self.page}',
                'class': f'page page-{self.page}',
                'ref': f'page-{self.page}',
                'style': 'width: 975px',
                'children': []
            }
        )

    def __addNewTable(self, root, borders=True):
        style = 'border-spacing: 0px;'
        root['children'].append(
            {
                'type': 'table',
                'id': f'{root["id"]} table-{self.table}',
                'class': f'table table-{self.table}' + ' table-bordered' if borders else '',
                'ref': f'{root["id"]} table-{self.table}',
                'style': style,
                'children': []
            }
        )

        self.t_row = 1

    def __addNewRow(self, root, block):
        style = ''

        root['children'].append(
            {
                'type': 'row',
                'id': f'{root["id"]} row-{self.t_row}',
                'class': f'row row-{self.t_row}',
                'ref': f'{root["id"]} row-{self.t_row}',
                'style': style,
                'children': []
            }
        )

        self.t_cell = 1

    def __addNewCell(self, root, block, borders=True):
        style = 'border: 1px solid black;' if borders else ''
        colspan = 1

        if  '<w:gridSpan w:val' in block._element.xml:
            colspan = int(re.search(r'(?<=<w:gridSpan w:val=")[0-9\.]+', block._element.xml)[0])

        # calc cell width in pt
        if block.width:
            width = block.width / 12700

            style += f'width: {width}pt;'

        root['children'].append(
            {
                'type': 'cell',
                'id': f'{root["id"]} cell-{self.t_cell}',
                'class': f'cell cell-{self.t_cell}',
                'ref': f'{root["id"]} cell-{self.t_cell}',
                'colspan': colspan,
                'rowspan': 1,
                'style': style,
                'children': []
            }
        )

    def __addNewUlToParent(self, root):
        root['children'].append(
            {
                'type': 'ul',
                'id': f'{root["id"]} ul-{self.ul}',
                'class': f'ul ul-{self.ul}',
                'ref': f'{root["id"]} ul-{self.ul}',
                'style': '',
                'children': []
            }
        )

    def __addBrToParent(self, root):
        root['children'].append(
            {
                'type': 'br',
                'id': f'{root["id"]} br-{self.br}',
                'class': f'br br-{self.br}',
                'ref': f'{root["id"]} br-{self.br}',
                'style': '',
                'children': []
            }
        )

        self.br += 1

    def __add_heading(self, root, block, type):
        style = self.__get_align(block)
        style += self.__get_font_size(block)
        style += self.__get_vertical_spacing(block)

        root['children'].append(
            {
                'type': f'h{type}',
                'id': f'{root["id"]} header-{self.header}',
                'class': f'h{type} h{type}-{self.header}',
                'ref': f'{root["id"]} header-{self.header}',
                'style': style,
                'text': block.text
            }
        )

        self.header += 1

    def __addImage(self, root, rIdx,  image):
        millis = int(round(time.time() * 1000))
        file_name = self.file_name.split('.')[0] + "-" + rIdx + "-" + str(millis) + ".png"
        fr = open(self.image_path + "/" + file_name, "wb")
        fr.write(image)
        fr.close()

        # encoded_string = base64.b64encode(image)
        encoded_string = b64encode(image).decode('utf-8')
        root['children'].append(
            {
                'type': 'img',
                'id': f'{root["id"]} img-{self.img}',
                'class': f'img img-responsive img-{self.img}',
                'ref': f'{root["id"]} img-{self.img}',
                'style': '',
                'image': encoded_string
            }
        )

        self.img += 1

    def __add_paragraph(self, root, block):
        style = self.__get_align(block)
        style += self.__get_font_size(block)
        style += self.__get_vertical_spacing(block)

        root['children'].append(
            {
                'type': 'p',
                'id': f'{root["id"]} paragraph-{self.paragraph}',
                'class': f'paragraph paragraph-{self.paragraph}',
                'ref': f'{root["id"]} paragraph-{self.paragraph}',
                'style': style,
                'text': block.text
            }
        )

        self.paragraph += 1

    def __get_align(self, block):
        style = ''
        if 'w:jc w:val="center"' in block._element.xml:
            style = 'text-align: center;'

        elif 'w:jc w:val="left"' in block._element.xml:
            style = 'text-align: left;'
            
        elif 'w:jc w:val="right"' in block._element.xml:
            style = 'text-align: right;'
            
        elif 'w:jc w:val="both"' in block._element.xml:
            style = 'text-align: justify;'

        return style

    def __get_vertical_spacing(self, block):
        """ вертикальные отступы до и после параграфа, заголовка, списка """

        style = ''

        spacing_line = re.search(r'<w:spacing [a-z0-9\.\:\=\" ]+/>', block._element.xml)

        if spacing_line:
            before = re.search(r'(?<=w:before=")[0-9\.]+', spacing_line[0])
            after = re.search(r'(?<=w:after=")[0-9\.]+', spacing_line[0])

            if before:
                style += f'margin-top: {int(before[0])/20}pt;'

            if after:
                style += f'margin-bottom: {int(after[0])/20}pt;'

        return style

    def __get_indentation(self, block):
        """ горизонтальные отступы до и после параграфа, заголовка, списка, красная строка """

        style = ''

        ind_line = re.search(r'<w:ind [a-z0-9\.\:\=\" ]+/>', block._element.xml)

        if ind_line:
            start = re.search(r'(?<=w:start=")[0-9\.]+', ind_line[0])
            end = re.search(r'(?<=w:end=")[0-9\.]+', ind_line[0])
            first_line = re.search(r'(?<=w:firstLine=")[0-9\.]+', ind_line[0]) or re.search(r'(?<=w:hanging=")[0-9\.]+', ind_line[0])

            if start:
                style += f'margin-left: {int(start[0])/12700}pt;'

            if end:
                style += f'margin-right: {int(end[0])/12700}pt;'

            if first_line:
                style += f'text-indent: {int(first_line[0])/12700}pt;'

        return style


    def __get_font_size(self, block):
        font_size = None

        if block.style and block.style.font and block.style.font.size:
            font_size = block.style.font.size / 12700.0

        elif re.search('(?<=<w:sz w:val=\")[0-9]+', block._element.xml):
            font_size = int(re.search('(?<=<w:sz w:val=\")[0-9]+', block._element.xml)[0]) / 2


        return f'font-size: {font_size}pt;' if font_size else ''

if __name__ == '__main__':

    import pprint
    import json

    pp = pprint.PrettyPrinter(indent=4)

    converter = Docx2HtmlConverter()

    # converter.convert('./', 'file-sample_100kB.docx')
    converter.convert('./', 'init_document.docx')

    # pp.pprint(converter.getJSON())
    # doc = Document('/home/ghost/design/programming/python/geekbrains/group_project/db_project_metrology_expertise/files/init_document.docx')
    # with open('xml.txt', 'w', encoding='utf8') as f:
    #         f.write(doc._element.xml)
    with open('json.txt', 'w', encoding='utf8') as f:
            f.write(json.dumps(converter.getJSON(), indent=4))
    
    

    
    